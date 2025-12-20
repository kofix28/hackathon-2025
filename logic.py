"""
Backend Logic for FieldScribe
Handles translation, image compression, and Word document generation
"""

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from deep_translator import GoogleTranslator
from datetime import datetime, date
from calendar import monthrange
from PIL import Image


def compress_image(image_file, max_width=800):
    """
    Takes a huge image file, resizes it, and returns a compressed byte stream.
    Handles PNG transparency to prevent crashes.
    """
    try:
        # CRITICAL FIX: Reset file pointer to the beginning.
        # This fixes the issue where camera images appear blank or missing.
        image_file.seek(0)

        image = Image.open(image_file)

        # 1. FIX PNG CRASH: Convert RGBA (Transparent) to RGB (Solid)
        if image.mode in ("RGBA", "P"):
            image = image.convert("RGB")

        # 2. Resize
        width_percent = (max_width / float(image.size[0]))
        new_height = int((float(image.size[1]) * float(width_percent)))
        image = image.resize((max_width, new_height), Image.Resampling.LANCZOS)

        # 3. Save
        img_byte_arr = BytesIO()
        image.save(img_byte_arr, format='JPEG', quality=70, optimize=True)
        img_byte_arr.seek(0)
        return img_byte_arr
    except Exception as e:
        print(f"Image compression error: {e}")
        return None


def set_paragraph_rtl_bidi(paragraph):
    """
    Sets paragraph alignment to right and adds bidi property for Hebrew RTL support.
    """
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Add bidi element
    pPr = paragraph._element.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)


def process_report(client_name, general_notes, defect_list, should_translate, report_mode='standard', logo_file=None):
    """
    Generates the Word Doc with professional card-based layout.
    """
    doc = Document()

    # Get page dimensions for full-width logo
    section = doc.sections[0]
    content_width = section.page_width - section.left_margin - section.right_margin

    # Add logo at the top of the document if provided
    if logo_file:
        logo_paragraph = doc.add_paragraph()
        set_paragraph_rtl_bidi(logo_paragraph)
        compressed_logo = compress_image(logo_file, max_width=600)  # Higher quality for full width
        if compressed_logo:
            run = logo_paragraph.add_run()
            run.add_picture(compressed_logo, width=content_width)  # Full content width
        doc.add_paragraph()  # Add space after logo

    # --- HELPER: Translator ---
    def t(text):
        if should_translate and text:
            try:
                return GoogleTranslator(source='auto', target='he').translate(text)
            except:
                return text
        return text

    # --- STYLE SETUP ---
    try:
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
    except:
        pass

    # --- HEADER & FOOTER ---
    section = doc.sections[0]
    header = section.header
    htable = header.add_table(1, 2, width=Inches(6))
    htable.autofit = False
    htable.columns[0].width = Inches(4)
    htable.columns[1].width = Inches(2)
    htable.cell(0, 0).text = f"Project ID: {client_name} | Engineer: איסמאיל ראבי"
    htable.cell(0, 1).text = datetime.now().strftime("%Y-%m-%d")

    footer = section.footer
    ftable = footer.add_table(1, 1, width=Inches(6))
    ftable.cell(0, 0).text = "Page "
    # Add page number field
    run = ftable.cell(0, 0).paragraphs[0].add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    run = ftable.cell(0, 0).paragraphs[0].add_run()
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    run._r.append(instrText)
    run = ftable.cell(0, 0).paragraphs[0].add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    run = ftable.cell(0, 0).paragraphs[0].add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)
    ftable.cell(0, 0).paragraphs[0].add_run(" of ")
    run = ftable.cell(0, 0).paragraphs[0].add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)
    run = ftable.cell(0, 0).paragraphs[0].add_run()
    instrText = OxmlElement('w:instrText')
    instrText.text = "NUMPAGES"
    run._r.append(instrText)
    run = ftable.cell(0, 0).paragraphs[0].add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar2)
    run = ftable.cell(0, 0).paragraphs[0].add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar3)
    # Set footer paragraph RTL and bidi
    set_paragraph_rtl_bidi(ftable.cell(0, 0).paragraphs[0])

    # --- TITLE PAGE ---
    meta_table = doc.add_table(rows=1, cols=2)
    meta_cell = meta_table.rows[0].cells[0]
    # Special design for client name
    run1 = meta_cell.paragraphs[0].add_run(t("שם הלקוח: "))
    run1.bold = True
    run1.font.size = Pt(12)
    run2 = meta_cell.paragraphs[0].add_run(client_name)
    run2.font.size = Pt(12)
    set_paragraph_rtl_bidi(meta_cell.paragraphs[0])
    doc.add_paragraph()
    title_text = f"DEFENSIVE OPINION: {client_name.upper()}" if report_mode == 'defensive' else f" : להלן חוות דעתי"
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(t(title_text))
    title_run.bold = True
    title_run.underline = True
    title_run.font.size = Pt(16)  # Larger size for title appearance
    set_paragraph_rtl_bidi(title_paragraph)

    

 

  

    # --- 2. DETAILED FINDINGS ---
    heading2 = doc.add_heading(t(" ממצאים מפורטים"), level=1)
    set_paragraph_rtl_bidi(heading2)

    image_counter = 1

    if defect_list:
        for defect in defect_list:
            # 1. Yellow Highlight Box
            yellow_table = doc.add_table(rows=1, cols=1)
            yellow_table.autofit = False
            yellow_table.columns[0].width = Inches(6)
            cell = yellow_table.cell(0, 0)
            cell.text = t(defect.get('title', 'Defect'))
            # Yellow background
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'FFFF00')
            cell._element.get_or_add_tcPr().append(shading)
            # Bold, 14pt, centered
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(14)
            set_paragraph_rtl_bidi(cell.paragraphs[0])

            # 2. Problem Definition
            desc = defect.get('desc', '')
            if desc:
                p = doc.add_paragraph(t(desc))
                set_paragraph_rtl_bidi(p)
            else:
                for _ in range(3):
                    p = doc.add_paragraph("")
                    set_paragraph_rtl_bidi(p)

            # 3. Evidence Grid
            photos = defect.get('photos', [])
            if not photos and 'photo' in defect:
                photos = [defect['photo']]
            if photos:
                # Create table with 2 columns
                num_rows = (len(photos) + 1) // 2
                evidence_table = doc.add_table(rows=num_rows, cols=2)
                evidence_table.style = 'Table Grid'  # Invisible borders? Actually, set to none
                # To make invisible, perhaps no style or custom
                for i, photo_file in enumerate(photos):
                    row = i // 2
                    col = 1 - (i % 2)  # RTL: Photo 1 right, Photo 2 left
                    cell = evidence_table.cell(row, col)
                    compressed = compress_image(photo_file)
                    if compressed:
                        run = cell.paragraphs[0].add_run()
                        run.add_picture(compressed, width=Inches(3))  # Half page width approx
            else:
                # Fallback: gray dashed box
                fallback_table = doc.add_table(rows=1, cols=1)
                fallback_table.autofit = False
                fallback_table.columns[0].width = Inches(6)
                cell = fallback_table.cell(0, 0)
                cell.text = 'הדבק תמונה כאן'
                set_paragraph_rtl_bidi(cell.paragraphs[0])
                # Gray background, dashed border
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'D3D3D3')  # Light gray
                cell._element.get_or_add_tcPr().append(shading)
                # Dashed border
                tcPr = cell._element.get_or_add_tcPr()
                borders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'dashed')
                    border.set(qn('w:sz'), '4')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), '000000')
                    borders.append(border)
                tcPr.append(borders)

            # 4. Standard Field
            p = doc.add_paragraph()
            standard = defect.get('code', '')
            if standard:
                p.add_run(t(standard))
            else:
                p.add_run('____________________')
            set_paragraph_rtl_bidi(p)

            # Large spacing between defects
            doc.add_page_break()

    else:
        p = doc.add_paragraph(t("No items recorded."))
        set_paragraph_rtl_bidi(p)

    
   # --- 1. EXECUTIVE SUMMARY ---
    heading1 = doc.add_heading(t(" הערות : "), level=3)
    set_paragraph_rtl_bidi(heading1)
    if general_notes:
        p = doc.add_paragraph(t(general_notes))
        set_paragraph_rtl_bidi(p)
    else:
        p = doc.add_paragraph(t("No specific general notes provided."))
        set_paragraph_rtl_bidi(p)
    # --- ENGINEER’S SIGN-OFF ---
    heading3 = doc.add_heading(t("המהנדס העורך והחותם: "), level=2)
    set_paragraph_rtl_bidi(heading3)
    p = doc.add_paragraph(t("חתימה: ___________________________"))
    set_paragraph_rtl_bidi(p)
    

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def get_calendar_month_data(year=None, month=None):
    """Returns calendar data for the CRM."""
    today = date.today()
    if year is None: year = today.year
    if month is None: month = today.month

    first_day_weekday, num_days = monthrange(year, month)
    first_day_weekday = (first_day_weekday + 1) % 7

    days = []
    for day in range(1, num_days + 1):
        day_date = date(year, month, day)
        days.append({
            'day': day,
            'date': day_date.strftime('%Y-%m-%d'),
            'weekday': day_date.weekday()
        })

    month_name = date(year, month, 1).strftime('%B')

    return {
        'days': days,
        'month_name': month_name,
        'year': year,
        'month': month,
        'first_day_weekday': first_day_weekday,
        'num_days': num_days
    }


def format_event_time(time_str):
    """Formats a time string (HH:MM) to a readable format (H:MM AM/PM)."""
    try:
        # Parse time string (format: "HH:MM")
        hour, minute = map(int, time_str.split(':'))
        
        # Convert to 12-hour format
        if hour == 0:
            return f"12:{minute:02d} AM"
        elif hour < 12:
            return f"{hour}:{minute:02d} AM"
        elif hour == 12:
            return f"12:{minute:02d} PM"
        else:
            return f"{hour - 12}:{minute:02d} PM"
    except Exception:
        # Return original string if parsing fails
        return time_str